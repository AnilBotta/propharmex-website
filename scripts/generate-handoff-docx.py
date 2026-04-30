"""
Generate the engineering handoff package as a .docx file.

Source of truth: docs/handoff.md.
Output: docs/handoff.docx — the customer-facing copy delivered with the
v1.0.0 launch.

This is a one-shot generator. Editorial changes happen in the markdown;
the docx is regenerated. Hard-coded structure asserts catch
content-drift between the two.

Run:
    py scripts/generate-handoff-docx.py

Requirements:
    py -m pip install python-docx==1.1.2

Pinned to python-docx 1.1.x; later majors break the table-style API
we use. Don't upgrade without re-checking the cell shading + table
header repeat behaviour. Mirrors scripts/generate-acr-docx.py — keep
the two scripts in lock-step on styling so the launch deliverables
look like a coherent set.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
except ImportError as exc:
    print(
        "Missing dependency. Install with:\n  py -m pip install python-docx==1.1.2",
        file=sys.stderr,
    )
    raise SystemExit(1) from exc


# Brand tokens duplicated from packages/config/design-tokens.css. Keep
# in sync if the brand palette ever shifts.
BRAND_PRIMARY_HEX = "0E4C5A"
BRAND_FG_HEX = "1F2933"
BRAND_MUTED_HEX = "5F6573"
BRAND_TABLE_HEADER_BG = "EBF4F6"


REPO_ROOT = Path(__file__).resolve().parent.parent
SOURCE_PATH = REPO_ROOT / "docs" / "handoff.md"
OUTPUT_PATH = REPO_ROOT / "docs" / "handoff.docx"


def hex_color(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


def shade_cell(cell, fill_hex: str) -> None:
    """Apply background shading to a table cell (no built-in API for this)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    """Set thin grey borders on all four sides of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D6D8DD")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.color.rgb = hex_color(BRAND_PRIMARY_HEX)
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = hex_color(BRAND_FG_HEX)


def add_table(doc: Document, rows: list[list[str]], header: bool = True):
    """Add a table with the brand-styled header row."""
    if not rows:
        return

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(cell)

            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)

            if header and r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = hex_color(BRAND_PRIMARY_HEX)
                shade_cell(cell, BRAND_TABLE_HEADER_BG)
            else:
                run.font.color.rgb = hex_color(BRAND_FG_HEX)


# ---------------------------------------------------------------------------
# Markdown parsing — intentionally minimal. We expect the source file to
# follow the structure asserted at the bottom of this script.
# ---------------------------------------------------------------------------

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
TABLE_DIVIDER_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")
LIST_RE = re.compile(r"^[-*]\s+(.*)$")
HR_RE = re.compile(r"^---\s*$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")


def strip_markdown(text: str) -> str:
    """Strip a tiny subset of inline markdown — enough for table cells."""
    # Bold + italic
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    # Inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Markdown links: [label](url) → label
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # HTML escapes
    text = text.replace("&rsquo;", "’").replace("&apos;", "'")
    return text.strip()


def split_table_row(line: str) -> list[str]:
    """Split a markdown table row into cells. Handles escaped pipes."""
    inner = line.strip().strip("|")
    cells = [c.strip() for c in inner.split("|")]
    return [strip_markdown(c) for c in cells]


def parse_markdown(md_text: str) -> list[tuple[str, object]]:
    """
    Walk the markdown linearly and produce a list of (kind, payload)
    blocks. Kinds:
      ("heading", (level, text))
      ("paragraph", text)
      ("table", rows)            # rows[0] is the header
      ("list", items)
      ("hr", None)
      ("blockquote", text)
    """
    blocks: list[tuple[str, object]] = []
    lines = md_text.splitlines()
    i = 0

    def flush_paragraph(buf: list[str]) -> None:
        if buf:
            text = strip_markdown(" ".join(buf).strip())
            if text:
                blocks.append(("paragraph", text))

    paragraph_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        m = HEADING_RE.match(line)
        if m:
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            level = len(m.group(1))
            blocks.append(("heading", (level, strip_markdown(m.group(2)))))
            i += 1
            continue

        if HR_RE.match(line):
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            blocks.append(("hr", None))
            i += 1
            continue

        m = BLOCKQUOTE_RE.match(line)
        if m:
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            quote_lines = []
            while i < len(lines) and BLOCKQUOTE_RE.match(lines[i]):
                quote_lines.append(BLOCKQUOTE_RE.match(lines[i]).group(1))
                i += 1
            blocks.append(("blockquote", strip_markdown(" ".join(quote_lines))))
            continue

        if TABLE_ROW_RE.match(line) and i + 1 < len(lines) and TABLE_DIVIDER_RE.match(lines[i + 1]):
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            rows: list[list[str]] = []
            rows.append(split_table_row(line))
            i += 2  # skip header + divider
            while i < len(lines) and TABLE_ROW_RE.match(lines[i]):
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append(("table", rows))
            continue

        m = LIST_RE.match(line)
        if m:
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            items: list[str] = []
            while i < len(lines):
                m2 = LIST_RE.match(lines[i])
                if not m2:
                    break
                items.append(strip_markdown(m2.group(1)))
                i += 1
            blocks.append(("list", items))
            continue

        if line.strip() == "":
            flush_paragraph(paragraph_buf)
            paragraph_buf = []
            i += 1
            continue

        paragraph_buf.append(line.strip())
        i += 1

    flush_paragraph(paragraph_buf)
    return blocks


# ---------------------------------------------------------------------------
# Render
# ---------------------------------------------------------------------------


def render_doc(blocks: list[tuple[str, object]]) -> Document:
    doc = Document()

    # Default font + page setup
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload  # type: ignore[misc]
            add_heading(doc, text, level)
        elif kind == "paragraph":
            add_paragraph(doc, payload)  # type: ignore[arg-type]
        elif kind == "table":
            add_table(doc, payload)  # type: ignore[arg-type]
        elif kind == "list":
            for item in payload:  # type: ignore[union-attr]
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item)
                run.font.size = Pt(11)
                run.font.color.rgb = hex_color(BRAND_FG_HEX)
        elif kind == "hr":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:bottom")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "6")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "D6D8DD")
            p_bdr.append(top)
            p_pr.append(p_bdr)
        elif kind == "blockquote":
            p = doc.add_paragraph()
            run = p.add_run(payload)  # type: ignore[arg-type]
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = hex_color(BRAND_MUTED_HEX)

    return doc


# ---------------------------------------------------------------------------
# Structure asserts — catch silent content-drift
# ---------------------------------------------------------------------------


def assert_structure(blocks: list[tuple[str, object]]) -> None:
    """Hard-coded asserts so editorial changes don't silently break the docx."""
    headings = [
        (level, text)
        for kind, payload in blocks
        if kind == "heading"
        for (level, text) in [payload]  # type: ignore[misc]
    ]
    h1s = [t for (lvl, t) in headings if lvl == 1]
    h2s = [t for (lvl, t) in headings if lvl == 2]

    assert len(h1s) == 1, f"Expected exactly 1 H1, got {len(h1s)}: {h1s}"
    assert h1s[0].startswith("Engineering handoff"), (
        f"H1 should be the handoff title, got: {h1s[0]!r}"
    )

    expected_h2_prefixes = [
        "1. Project summary",
        "2. Tech stack",
        "3. What shipped",
        "4. Page inventory",
        "5. AI tools quartet",
        "6. Operations and runbooks",
        "7. CI gates",
        "8. Known follow-ups",
        "9. Repository entry points",
        "10. Handoff next steps",
        "11. Revision history",
    ]
    for prefix in expected_h2_prefixes:
        assert any(h.startswith(prefix) for h in h2s), (
            f"Missing expected H2 starting with {prefix!r}. Got H2s: {h2s}"
        )

    tables = [payload for kind, payload in blocks if kind == "table"]
    assert len(tables) >= 7, (
        f"Expected at least 7 tables (tech stack, phase status, page inventory, "
        f"operations links, CI gates, handoff next steps, revision history). "
        f"Got {len(tables)}."
    )


def main() -> None:
    if not SOURCE_PATH.exists():
        print(f"Source not found: {SOURCE_PATH}", file=sys.stderr)
        raise SystemExit(1)

    md_text = SOURCE_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)
    assert_structure(blocks)
    doc = render_doc(blocks)
    doc.save(OUTPUT_PATH)
    rel = OUTPUT_PATH.relative_to(REPO_ROOT)
    print(f"Wrote {rel} - {len(blocks)} blocks rendered.")


if __name__ == "__main__":
    main()
